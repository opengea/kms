from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new document
extended_doc = Document()

# Define styles
styles = extended_doc.styles

# Title style
title_style = styles.add_style('TitleStyle', 1)
title_font = title_style.font
title_font.name = 'Arial'
title_font.size = Pt(24)
title_font.bold = True

# Heading 1 style
heading1_style = styles.add_style('Heading1Style', 1)
heading1_font = heading1_style.font
heading1_font.name = 'Arial'
heading1_font.size = Pt(20)
heading1_font.bold = True

# Heading 2 style
heading2_style = styles.add_style('Heading2Style', 1)
heading2_font = heading2_style.font
heading2_font.name = 'Arial'
heading2_font.size = Pt(18)
heading2_font.bold = True

# Normal text style
normal_style = styles.add_style('NormalStyle', 1)
normal_font = normal_style.font
normal_font.name = 'Arial'
normal_font.size = Pt(12)

# Add title
title = extended_doc.add_paragraph('DANI MMO-CLIR Thematic', style='TitleStyle')
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Introduction
extended_doc.add_paragraph('Introduction', style='Heading1Style')
intro = extended_doc.add_paragraph(style='NormalStyle')
intro.add_run(
    "DANI MMO-CLIR Thematic is a proprietary system designed to solve complex multilingual information retrieval problems. "
    "It allows users to query information in multiple languages and retrieve results in their preferred language, making it "
    "an invaluable tool for diverse and international data needs. This document outlines the key features, modules, technical "
    "characteristics, and scope of the DANI MMO-CLIR Thematic system.\n\n"
    "In today's globalized world, information retrieval across languages is a significant challenge. Traditional search engines "
    "often fail to provide accurate translations and contextual relevance, leading to inefficiencies. DANI MMO-CLIR Thematic addresses "
    "these issues by leveraging advanced semantic analysis and high-performance computing to deliver precise and relevant search results."
)

# Product Overview
extended_doc.add_paragraph('Product Overview', style='Heading1Style')
overview = extended_doc.add_paragraph(style='NormalStyle')
overview.add_run(
    "The DANI MMO-CLIR Thematic system integrates Multilingual, Multi-Object, and Cross-Language Information Retrieval capabilities. "
    "It processes not only web content but also digital documents such as PDFs and office files, indexing them for efficient retrieval. "
    "The system is particularly tailored to the automotive industry, covering a wide range of related fields including metallurgy, electronics, and design.\n\n"
    "DANI MMO-CLIR stands for Multilingual Multi-Object Cross-Language Information Retrieval. This innovative system is designed to "
    "facilitate information discovery and retrieval across different languages and formats. By indexing both textual and non-textual content, "
    "the system ensures comprehensive search results that are contextually relevant and linguistically accurate."
)

# Key Features
extended_doc.add_paragraph('Key Features', style='Heading1Style')
features = extended_doc.add_paragraph(style='NormalStyle')
features.add_run("The main features of the DANI MMO-CLIR Thematic system include:\n").bold = True
features.add_run(
    "- Multilingual Information Retrieval: Allows queries in multiple languages and returns results translated into the user's preferred language.\n"
    "- Multi-Object Processing: Indexes various digital documents, making them searchable alongside web content.\n"
    "- Cross-Language Retrieval: Retrieves relevant documents in different languages based on the user's query language.\n"
    "- Thematic Focus: Specializes in the automotive industry with peripheral connections to related fields.\n"
    "- Real-Time Querying: Provides fast, real-time responses to user queries.\n\n"
    "These features make DANI MMO-CLIR Thematic a powerful tool for businesses and researchers who need to navigate and analyze large volumes of "
    "multilingual information efficiently. The system's ability to handle complex queries and provide accurate translations ensures that users can "
    "access the information they need, regardless of language barriers."
)

# Technical Characteristics
extended_doc.add_paragraph('Technical Characteristics', style='Heading1Style')
tech_char = extended_doc.add_paragraph(style='NormalStyle')
tech_char.add_run(
    "The DANI MMO-CLIR Thematic system leverages high-performance computing (HPC) to handle extensive data processing tasks. "
    "It requires GPU-accelerated servers with the following specifications:\n"
    "- Dual Intel Xeon processors\n"
    "- 128 GB RAM\n"
    "- 2 TB of expanded memory\n"
    "- Multiple SATA-3 hard drives\n\n"
    "The system initially supports Catalan, Spanish, French, English, and German, with future expansions planned for additional languages.\n\n"
    "High-performance computing enables the system to process vast amounts of data quickly and accurately. By utilizing GPU-accelerated servers, "
    "DANI MMO-CLIR Thematic can perform complex semantic analysis and multilingual translations in real-time, ensuring that users receive timely and relevant search results."
)

# System Modules
extended_doc.add_paragraph('System Modules', style='Heading1Style')
modules = extended_doc.add_paragraph(style='NormalStyle')
modules.add_run("The DANI MMO-CLIR Thematic system is divided into two main modules:\n").bold = True

extended_doc.add_paragraph('1. DANI_CREA', style='Heading2Style')
mod1 = extended_doc.add_paragraph(style='NormalStyle')
mod1.add_run(
    "This module generates a comprehensive catalog of indexed digital documents and web content. "
    "Each catalog entry includes demographic data and inferred activities based on semantic analysis. "
    "The catalog is designed for efficient multilingual querying and retrieval.\n\n"
    "DANI_CREA automates the process of cataloging vast amounts of digital content. By performing semantic and morphological analysis, "
    "the system can deduce the activities and subjects contained within each document. This results in a rich, searchable database that "
    "provides users with a detailed overview of the available information."
)

extended_doc.add_paragraph('2. DANI_WEB', style='Heading2Style')
mod2 = extended_doc.add_paragraph(style='NormalStyle')
mod2.add_run(
    "The web module provides a functional web interface for querying the catalog in real-time. "
    "It supports multilingual queries and displays results in a user-friendly format, allowing users to filter, group, and export data as needed.\n\n"
    "DANI_WEB ensures that users can access the indexed data from any web browser. The interface is designed to be intuitive and responsive, "
    "providing a seamless user experience. Advanced filtering and grouping options enable users to refine their searches and extract the most relevant information."
)

# Implementation and Use Cases
extended_doc.add_paragraph('Implementation and Use Cases', style='Heading1Style')
implementation = extended_doc.add_paragraph(style='NormalStyle')
implementation.add_run(
    "Implementing DANI MMO-CLIR Thematic involves setting up the required hardware and configuring the system according to the client's needs. "
    "The process includes:\n"
    "- Installing the necessary servers and hardware components\n"
    "- Configuring the software and database\n"
    "- Importing and indexing existing digital content\n"
    "- Training users on how to utilize the system effectively\n\n"
    "Once implemented, the system can be used in various scenarios, such as:\n"
    "- Research and academic studies: Facilitating multilingual literature reviews and data analysis.\n"
    "- Business intelligence: Analyzing market trends and competitor information across different languages.\n"
    "- Legal research: Accessing multilingual legal documents and case studies.\n\n"
    "By tailoring the implementation process to the specific needs of each client, DANI MMO-CLIR Thematic ensures that users can leverage its full capabilities for their unique requirements."
)

# Scope and Limitations
extended_doc.add_paragraph('Scope and Limitations', style='Heading1Style')
scope = extended_doc.add_paragraph(style='NormalStyle')
scope.add_run(
    "DANI MMO-CLIR Thematic is proprietary software owned by Frederic Monràs i Vidiella and Miquel Monràs i Ho. "
    "It is not open source, but it complies with international standards for data import and export. "
    "The system does not include hardware; clients must provide their own servers if they wish to host the system locally. "
    "The dictionaries and ontologies used in the system are extensive but not exhaustive, and may not cover 100% of the terminology in some cases.\n\n"
    "The proprietary nature of the software ensures that clients receive a tailored and optimized solution. However, this also means that "
    "clients need to be aware of the limitations regarding hardware requirements and the scope of the included dictionaries and ontologies."
)

# Example Use Case
extended_doc.add_paragraph('Example Use Case', style='Heading1Style')
example = extended_doc.add_paragraph(style='NormalStyle')
example.add_run(
    "Consider a user querying the system for information on 'automobiles'. The system will:\n"
    "- Expand the query to include synonyms and related terms in multiple languages.\n"
    "- Retrieve and display relevant documents from the indexed catalog.\n"
    "- Allow the user to filter and export the results in various formats, such as CSV or Excel.\n\n"
    "In this scenario, the user starts by entering the search term 'automobiles'. The system performs a comprehensive search, "
    "identifying related terms such as 'cars', 'vehicles', and 'automóviles' across different languages. The results are displayed "
    "in a table format, highlighting the most relevant entries. Users can further refine their search by applying filters and grouping the results by various criteria."
)

# Save the final extended document
extended_doc_path = "/mnt/data/Extended_Creative_DANI_MMO-CLIR_Thematic_Final.docx"
extended_doc.save(extended_doc_path)

extended_doc_path &#8203;:citation[oaicite:0]{index=0}&#8203;

